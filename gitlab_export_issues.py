#
# Gitlab Issues Export 2025/2/7
# pip install python-gitlab 
# pip install python-docx
##################
GITLAB_DOMAIN = "************" #your Gitlab domain
GITLAB_AUTH_TOKEN = '************' # https://YOUR_GITLAB/-/user_settings/personal_access_tokens
PROJETCT_ID = ** # Project ID
ISSUE_NUMBER_BIGGER_THAN = None
ISSUE_NUMBER_LESS_THAN = None
EXPORT_DOCX = 1 # 0:csv 1:docx
##################
"""Example script for showing example usage of python-gitlab package."""
import csv
import gitlab
import os
import re
import docx
from docx import Document

def build_csv_issue(filename, one_issue):
    if not one_issue:
        return

    # Extract iid, title, and description
    iid = one_issue.attributes['iid']
    if (ISSUE_NUMBER_BIGGER_THAN is not None) and (ISSUE_NUMBER_BIGGER_THAN > iid):
        return;
    if (ISSUE_NUMBER_LESS_THAN is not None) and (ISSUE_NUMBER_LESS_THAN < iid):
        return;

    title = one_issue.attributes['title']
    description = one_issue.attributes['description']
    state = one_issue.attributes['state']
    created_at = one_issue.attributes['created_at']
    updated_at = one_issue.attributes['updated_at']
    web_url = one_issue.attributes['web_url']

    notes = one_issue.notes.list(iterator=True)
    sorted_notes = sorted(notes, key=lambda note: note.created_at)
    for note in sorted_notes:
        #print(note.body)
        my_str = note.body or ""
        if (description is not None) and (my_str is not None):
            description += "\n------------------\n" + my_str

    #Replace multiple newlines with one
    #Replace multiple newlines with one
    new_description = ""
    if description is not None:
        new_description = re.sub(r'\n+', '\n', '\n'.join(line.rstrip() for line in description.split('\n')))

    my_num = str(iid)+"\n" or ""
    print(my_num)
    
    doc = docx.Document(filename)
    
    
    with open(filename, 'a', encoding='utf-8-sig', newline='') as file:  # Open in append mode ('a')
        writer = csv.writer(file)
        writer.writerow(["[ISSUE]"+str(iid), "[TITLE]"+title, "[DESCRIPTION]"+new_description, "[STATE]"+state, "[CREATED]"+created_at, "[UPDATED"+updated_at, "[URL]"+web_url])
        csv_file = filename
        pdf_file = filename+".pdf"


def build_docx_issue(doc, one_issue):
    if not one_issue:
        return

    # Extract iid, title, and description
    iid = one_issue.attributes['iid']
    if (ISSUE_NUMBER_BIGGER_THAN is not None) and (ISSUE_NUMBER_BIGGER_THAN > iid):
        return;
    if (ISSUE_NUMBER_LESS_THAN is not None) and (ISSUE_NUMBER_LESS_THAN < iid):
        return;

    title = one_issue.attributes['title']
    description = one_issue.attributes['description']
    state = one_issue.attributes['state']
    created_at = one_issue.attributes['created_at']
    updated_at = one_issue.attributes['updated_at']
    web_url = one_issue.attributes['web_url']

    notes = one_issue.notes.list(iterator=True)
    sorted_notes = sorted(notes, key=lambda note: note.created_at)
    for note in sorted_notes:
        #print(note.body)
        my_str = note.body or ""
        if (description is not None) and (my_str is not None):
            description += "\n------------------\n" + my_str

    #Replace multiple newlines with one
    #Replace multiple newlines with one
    new_description = ""
    if description is not None:
        new_description = re.sub(r'\n+', '\n', '\n'.join(line.rstrip() for line in description.split('\n')))
        new_description = re.sub(r'[\x00-\x1F\x7F]', '', new_description)

    my_num = str(iid)+"\n" or ""
    print(my_num)
    doc.add_paragraph("[ISSUE] "+str(iid))
    doc.add_paragraph("[TITLE] "+title)
    doc.add_paragraph("[CREATED] "+created_at)
    doc.add_paragraph("[UPDATED "+updated_at)
    doc.add_paragraph("[URL] "+web_url)
    doc.add_paragraph("[DESCRIPTION] "+new_description.encode('utf-8').decode('utf-8'))

def main():
    issue_counter = 0;

    # Establish GitLab connection and get project
    git = gitlab.Gitlab('https://' + GITLAB_DOMAIN, private_token=GITLAB_AUTH_TOKEN)
    git.auth()
    project = git.projects.get(PROJETCT_ID)

    doc = Document()

    issues = project.issues.list(iterator=True)
    for issue in issues:
        #print(issue)
        issue_counter += 1;
        filename_mod = int(issue_counter / 100);
        filename = "Gitlab_Issues_" + str(filename_mod).zfill(4) + ".docx"

        if EXPORT_DOCX == 1:
            build_docx_issue(doc, issue)
            if( int(issue_counter / 100) == issue_counter / 100):
                doc.save(filename)
                doc = Document()
        else:
            build_csv_issue(filename, issue)

    doc.save('Gitlab_Issues_last.docx')

if __name__ == "__main__":
    main()